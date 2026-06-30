from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "QSentia_Website_Technical_Report_IEEE_Aligned.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(23, 37, 84)
MUTED = RGBColor(92, 104, 95)
LIGHT_BLUE = "EEF2FF"
HEADER_FILL = "F2F4F7"
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)

TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line_spacing=1.10):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_para(doc, text="", style=None, size=11, color=BLACK, bold=False, italic=False, before=0, after=6):
    paragraph = doc.add_paragraph(style=style)
    set_paragraph_spacing(paragraph, before=before, after=after)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def set_cell_text(cell, text, bold=False, color=BLACK, size=10.5, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.05)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in CELL_MARGIN_DXA.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)


def set_table_borders(table, color="D9E0F2"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_table(doc, headers, rows, widths, header_fill=HEADER_FILL, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    header = table.rows[0]
    mark_header_row(header)
    for i, label in enumerate(headers):
        shade_cell(header.cells[i], header_fill)
        set_cell_text(header.cells[i], label, bold=True, color=NAVY, size=font_size)
    for row_values in rows:
        row = table.add_row()
        for i, value in enumerate(row_values):
            set_cell_text(row.cells[i], str(value), size=font_size)
    add_para(doc, "", after=2)
    return table


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(paragraph, before=0, after=4, line_spacing=1.167)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=BLACK)
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    set_paragraph_spacing(paragraph, before=0, after=4, line_spacing=1.167)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=BLACK)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    if level == 1:
        set_paragraph_spacing(paragraph, before=16, after=8)
        size, color = 16, BLUE
    elif level == 2:
        set_paragraph_spacing(paragraph, before=12, after=6)
        size, color = 13, BLUE
    else:
        set_paragraph_spacing(paragraph, before=8, after=4)
        size, color = 12, DARK_BLUE
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=True)
    return paragraph


def add_callout(doc, title, body, fill="EEF2FF"):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    set_table_borders(table, color="C7D2FE")
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_margins(cell)
    shade_cell(cell, fill)
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, before=0, after=3)
    run = paragraph.add_run(title)
    set_run_font(run, size=11, color=NAVY, bold=True)
    paragraph2 = cell.add_paragraph()
    set_paragraph_spacing(paragraph2, before=0, after=0)
    run2 = paragraph2.add_run(body)
    set_run_font(run2, size=10.5, color=BLACK)
    add_para(doc, "", after=2)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    set_paragraph_spacing(header, before=0, after=0)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("QSentia Investor Site - Technical Report")
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    set_paragraph_spacing(footer, before=0, after=0)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Prepared for internal review - IEEE-aligned documentation, not a certification claim")
    set_run_font(run, size=8.5, color=MUTED)


def add_title_page(doc):
    add_para(doc, "QSentia", size=12, color=MUTED, bold=True, after=6)
    title = doc.add_paragraph()
    set_paragraph_spacing(title, before=14, after=4)
    run = title.add_run("Website Technical Report")
    set_run_font(run, size=25, color=BLACK, bold=True)
    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, before=0, after=16)
    run = subtitle.add_run("Implementation wiring, API data flow, page inventory, loading-state behavior, and IEEE-aligned standards mapping")
    set_run_font(run, size=13.5, color=MUTED)

    rows = [
        ("Project", "QSentia Investor Intelligence Platform"),
        ("Repository", "QSentia-com/qsentia-investor-site"),
        ("Framework", "Next.js 16.2.4, React 19.2.4, TypeScript 5, Tailwind CSS 4"),
        ("Prepared", "June 4, 2026"),
        ("Status", "Implementation report based on current local workspace and live localhost API responses"),
    ]
    add_table(doc, ["Field", "Value"], rows, [1800, 7560], font_size=10)
    add_callout(
        doc,
        "Scope note",
        "This report documents the website as implemented in the current workspace. The IEEE section maps the implementation to relevant IEEE and ISO/IEC/IEEE software engineering standards; it does not assert formal third-party certification.",
    )
    doc.add_page_break()


def build_report():
    doc = Document()
    configure_document(doc)
    add_header_footer(doc)
    add_title_page(doc)

    add_heading(doc, "1. Executive Summary", 1)
    add_para(
        doc,
        "The QSentia investor website is implemented as a Next.js application that presents live model telemetry, research views, marketplace model profiles, MLEQ engine context, account access, legal pages, and an embedded assistant named Alex. The build was moved away from dummy or placeholder-looking data toward API-backed telemetry sourced from the application API layer and, for local development, upstream QSentia dashboard data.",
    )
    add_para(
        doc,
        "The main implementation objective was professionalizing the experience: replacing generic AI-looking copy and dummy cards, wiring pages to the API folder, adding professional loading buffers while slow telemetry resolves, explaining model equity indices in plain investor language, and aligning page content with mission, risk, and alpha narratives.",
    )
    add_table(
        doc,
        ["Area", "Current implementation evidence", "Result"],
        [
            ("Live telemetry", "/api/dashboard reports 10 registry models, 10 comparison rows, 5 benchmarks, 6 selected equity-curve points, 6 decisions, and 3 submitted orders.", "Pages consume source-backed values instead of mock data."),
            ("API marketplace", "/api/models maps dashboard modelComparison data into marketplace cards; detail pages use /api/models/{slug}.", "Model cards and profiles remain connected to dashboard telemetry."),
            ("Research terminal", "ResearchTerminal consumes /api/dashboard, fetches per-model detail curves when needed, and renders tickers, filters, tabs, and normalized charts.", "Live research view supports model and benchmark comparison."),
            ("Loading buffers", "ApiLoadingPanel is shared across dashboard, research, marketplace, model detail, and MLEQ.", "Slow API calls present a polished buffering state rather than premature empty states."),
        ],
        [1800, 5200, 2360],
        font_size=9,
    )

    add_heading(doc, "2. Technology Stack and Runtime Wiring", 1)
    add_table(
        doc,
        ["Layer", "Implementation", "Notes"],
        [
            ("Application framework", "Next.js 16 app router under app/", "Client pages use useSWR for API-backed state; route handlers live under app/api."),
            ("UI runtime", "React 19.2.4 with TypeScript", "Componentized page shell, navigation, footer, assistant, dashboard, research terminal, and loading panel."),
            ("Styling", "Tailwind CSS 4 utility classes", "Professional blue/navy palette, grid backdrops, compact cards, responsive layouts."),
            ("Charts", "Recharts and custom SVG", "Dashboard/research use Recharts; home and MLEQ use custom normalized equity visuals."),
            ("Data parsing and metrics", "PapaParse plus lib/metrics.ts", "CSV rows are normalized, converted to returns, drawdowns, normalized equity index, and risk statistics."),
            ("Data fetching", "SWR on client pages; fetch inside route handlers", "Refresh intervals are used for dashboard/research and assistant data."),
        ],
        [1900, 3100, 4360],
        font_size=9,
    )

    add_heading(doc, "3. Page and Component Inventory", 1)
    add_para(doc, "The site uses a shared PageShell and SiteHeader from components/PageChrome.tsx. Root layout injects SiteFooter and AlexAssistant globally so every route keeps footer links and assistant access.")
    add_table(
        doc,
        ["Route or component", "Primary responsibility", "Data wiring"],
        [
            ("/", "Investor-facing homepage, value proposition, registry/product cards, selected model equity index, benchmark comparison, live response preview.", "SWR GET /api/dashboard."),
            ("/dashboard", "Operating telemetry terminal with model selector, metrics, normalized curves, execution status, registry table, benchmark table, and raw source tables.", "SWR GET /api/dashboard or /api/dashboard?model={id}."),
            ("/research", "Research overview page that embeds ResearchTerminal.", "ResearchTerminal calls /api/dashboard and model-specific /api/dashboard?model={id} when curve detail is needed."),
            ("/mleq", "Dark-themed Machine Learning Equity Quant narrative page with live stats, model rank view, thesis, pipeline monitor, and discipline sections.", "SWR GET /api/dashboard."),
            ("/marketplace", "Model-product marketplace with search and category filters.", "SWR GET /api/models."),
            ("/marketplace/{slug}", "Single model profile, telemetry preview, source details, features, access CTA.", "SWR GET /api/models/{slug}; POST /api/models/{slug}/demo for preview."),
            ("/docs", "API endpoint documentation and curl examples.", "Static content aligned to implemented API routes."),
            ("/signin and /create-account", "Client-side account forms.", "Local storage session via lib/clientSession.ts."),
            ("AlexAssistant", "Help assistant named Alex, personalized by stored user name and aware of website pages and live telemetry.", "Fetches /api/dashboard and /api/models; reads local session."),
            ("SiteFooter", "Professional footer with docs, API, company/legal, and social links.", "Static link wiring."),
        ],
        [2100, 4100, 3160],
        font_size=8.5,
    )

    add_heading(doc, "4. API and Data-Flow Wiring", 1)
    add_heading(doc, "4.1 Dashboard API", 2)
    add_para(
        doc,
        "The central route is app/api/dashboard/route.ts. It is declared dynamic and uses the Node.js runtime. For localhost and 127.0.0.1, it attempts an upstream dashboard proxy to https://www.qsentia.com before falling back to local GitHub/raw-log assembly. It maintains a last-good cache under .qsentia-cache/dashboard-last-good.json to reduce dead-end states when upstream data is temporarily unavailable.",
    )
    add_table(
        doc,
        ["Input/source", "Processing", "Output fields"],
        [
            ("Required model registry and repository logs", "Fetches model configs and CSV/JSON logs; normalizes portfolio rows, decisions, orders, positions, target weights, health status, execution realism, and readiness checks.", "registry, selectedModelConfig, latest, decisions, positions, targetWeights, plannedOrders, submittedOrders."),
            ("Portfolio/account observations", "Builds daily portfolio values, computes selected performance values, returns, drawdowns, and normalized equity index using normalizeTo100.", "stats, equityCurve, returns, drawdowns, latest.portfolioValue, portfolioReturn."),
            ("Benchmarks", "Fetches benchmark series for SPY, QQQ, DIA, IWM, and VTI and computes performance statistics.", "benchmarks with points, stats, and rowCount."),
            ("Model comparison loop", "For every registered model, fetches portfolio/latest account rows, computes normalized curves and stats, and attaches model-specific benchmarks.", "modelComparison with id, name, repo, logsPath, points, stats, latestValue, row counts, inceptionDate."),
            ("Summary mode", "GET /api/dashboard?summary=1 returns registry-focused lightweight response.", "summary payload for live API ribbon and lightweight status surfaces."),
        ],
        [2400, 4100, 2860],
        font_size=8.5,
    )

    add_heading(doc, "4.2 Marketplace API", 2)
    add_para(doc, "The marketplace API is a thin layer over the dashboard. lib/modelCatalog.ts fetches /api/dashboard, maps modelComparison entries into marketplace models, generates slugs, categories, tags, and performance values, and uses model-specific dashboard calls for detail pages.")
    add_table(
        doc,
        ["Endpoint", "Method", "Implementation behavior"],
        [
            ("/api/models", "GET", "Returns model list from getLiveMarketplaceModels, source app-api-dashboard, and a timestamp. Cache-Control allows public s-maxage and stale-while-revalidate."),
            ("/api/models/{slug}", "GET", "Resolves slug to a model, fetches model-specific dashboard detail, returns long description, features, source details, latest state, and metrics."),
            ("/api/models/{slug}/demo", "POST", "Rate-limited to 5 requests per hour per IP. Returns latest decision preview fields: action, confidence, positionSize, latency, timestamp, and source."),
        ],
        [2500, 1000, 5860],
        font_size=9,
    )

    add_heading(doc, "4.3 Client Data Wiring", 2)
    add_table(
        doc,
        ["Client surface", "Hook/endpoint", "State behavior"],
        [
            ("Home", "useSWR('/api/dashboard')", "Shows homepage telemetry, model registry count, equity index explanation, and benchmark comparison."),
            ("Dashboard", "useSWR(endpoint)", "Endpoint changes when selected model changes. Initial load is buffered with ApiLoadingPanel."),
            ("ResearchTerminal", "useSWR('/api/dashboard') plus conditional model detail fetcher", "If modelComparison lacks curve points, it fetches per-model dashboard data and builds research curves."),
            ("MLEQ", "useSWR('/api/dashboard')", "Dark loading buffer appears before model families, pipeline, and terminal metrics render."),
            ("Marketplace", "useSWR('/api/models')", "Search and category filtering operate on API-backed model list."),
            ("Model detail", "useSWR('/api/models/{slug}')", "Initial load displays profile buffer; POST demo route is invoked manually by the preview button."),
            ("Alex assistant", "fetch('/api/dashboard') and fetch('/api/models')", "Assistant can summarize live telemetry and personalize greeting from local storage session."),
        ],
        [2100, 2650, 4610],
        font_size=8.5,
    )

    add_heading(doc, "5. UX Wiring and Professionalization Changes", 1)
    add_bullet(doc, "Homepage hero copy was changed to investor-native language: More alpha. Less unmanaged risk.")
    add_bullet(doc, "The previous pill-style AI-looking labels were removed or replaced with quieter institutional copy.")
    add_bullet(doc, "Selected model performance was renamed to Selected model equity index and explained as an index-style track record where the first observation equals 100.")
    add_bullet(doc, "The home chart now uses current index, source points, and start-100 terminology to avoid confusion with currency or Sensex/Nifty style market indices.")
    add_bullet(doc, "The shared ApiLoadingPanel prevents premature empty states while live dashboard/model data is still loading.")
    add_bullet(doc, "Footer, problem-solution, shipping-policy, mission/vision/objective content, and MLEQ pages were professionalized around source-backed telemetry and risk review.")
    add_bullet(doc, "Alex assistant is globally available and can reference website pages, API docs, live telemetry, model registry, account access, and MLEQ.")

    add_heading(doc, "6. Standards and Compliance Mapping", 1)
    add_callout(
        doc,
        "Interpretation",
        "The following matrix maps the implementation to relevant IEEE and ISO/IEC/IEEE standards. It is an engineering-alignment report, not a formal audit certificate.",
        fill="F8FAFF",
    )
    add_table(
        doc,
        ["Standard", "Purpose", "Implementation evidence", "Recommended next step"],
        [
            ("ISO/IEC/IEEE 42010", "Architecture description", "This report captures stakeholder concerns, system context, API views, data views, and route/component views.", "Maintain an architecture decision record for upstream proxy, cache, and telemetry source decisions."),
            ("IEEE 1016", "Software Design Description", "Pages, components, API endpoints, data stores, interfaces, and loading-state behavior are documented.", "Add sequence diagrams for dashboard and marketplace calls in future design docs."),
            ("ISO/IEC/IEEE 29148", "Requirements engineering and SRS", "Functional requirements are traceable to pages and endpoints: dashboard telemetry, marketplace models, model detail, demo preview, account/session personalization.", "Create a formal SRS with requirement IDs and acceptance criteria."),
            ("IEEE 830 legacy SRS guidance", "Legacy SRS structure still used by many reviewers", "Report sections mirror SRS concerns: purpose, scope, interfaces, data, constraints, quality attributes.", "Use only as legacy reference; prefer 29148 for future requirements."),
            ("IEEE 730", "Software quality assurance planning", "Linting, TypeScript checks, and browser smoke tests were used after implementation changes.", "Document a recurring QA plan with owners, environments, pass/fail thresholds, and release gates."),
            ("IEEE 828", "Configuration management", "Repo, environment variables, route files, and .qsentia-cache last-good behavior are identified.", "Create release tags and configuration baselines for production deployments."),
            ("IEEE 829 / ISO/IEC/IEEE 29119", "Test documentation and software testing", "Manual browser verification and command checks were performed for route loading, data rendering, and console errors.", "Formalize test cases for API failure, upstream fallback, empty data, and model switching."),
            ("IEEE 1063", "Software user documentation", "Docs page lists API endpoints and usage flow; page copy explains model equity index and live telemetry context.", "Add user-facing data dictionary for equity index, drawdown, Sharpe, and source rows."),
            ("ISO/IEC/IEEE 12207", "Software lifecycle processes", "Implementation, verification, API integration, and maintenance concerns are represented.", "Define maintenance procedures for registry updates, benchmark changes, and API schema changes."),
            ("ISO/IEC/IEEE 24765", "Systems/software vocabulary", "Report uses consistent terms: telemetry, model registry, normalized equity index, source rows, benchmark series, execution audit.", "Keep a terminology appendix in product docs."),
            ("IEEE 1028", "Software reviews and audits", "The document can support technical review of wiring, interface behavior, and standards alignment.", "Run a structured peer review before investor or compliance distribution."),
        ],
        [1900, 2100, 3450, 1910],
        font_size=7.8,
    )

    add_heading(doc, "7. Quality Attributes", 1)
    add_table(
        doc,
        ["Attribute", "Current design response", "Residual risk"],
        [
            ("Reliability", "Upstream dashboard attempts, last-good cache fallback, and controlled empty/loading states reduce visible failure.", "Local cache age and upstream schema drift need monitoring."),
            ("Data integrity", "No synthetic model metrics are introduced by client pages; missing values remain unavailable until source rows exist.", "Report readers should understand that live upstream values can change after the document is prepared."),
            ("Performance", "SWR caching and API cache headers reduce repeated fetch work; summary endpoint supports lightweight status surfaces.", "Dashboard route can be slow when fetching all model/benchmark data; background caching or batch endpoints could improve response time."),
            ("Usability", "Plain-language equity index explanation, professional loading buffers, model selector, filters, and tabs improve comprehension.", "More in-product definitions would help less technical investors."),
            ("Security/privacy", "GitHub token candidates are server-side in route handler; demo preview is rate-limited. Client account session uses local storage only.", "Local storage is not a secure authentication mechanism; production auth should use server-backed sessions."),
            ("Maintainability", "Shared PageShell, ApiLoadingPanel, metrics utilities, model catalog, and research terminal isolate repeated behavior.", "Legacy *-SHLOK files should be reviewed and archived or removed to reduce repository noise."),
        ],
        [1750, 4350, 3260],
        font_size=8.4,
    )

    add_heading(doc, "8. Verification Evidence", 1)
    add_para(doc, "The following checks were performed during implementation and report preparation.")
    add_table(
        doc,
        ["Check", "Command or method", "Result"],
        [
            ("TypeScript", "npx tsc --noEmit", "Passed after loading-state changes."),
            ("ESLint", "npx eslint on changed API-backed pages and shared components", "Passed after wiring updates."),
            ("Local API sample", "GET http://localhost:3001/api/dashboard", "Returned 10 registry models, 10 model comparison rows, 5 benchmarks, 6 selected curve points, 6 decisions, and 3 submitted orders."),
            ("Marketplace API sample", "GET http://localhost:3001/api/models", "Returned 10 models, source app-api-dashboard, and slug qsentia-brppo-macro-rotation-alpaca for the first model."),
            ("Browser smoke test", "In-app browser checks for /dashboard, /research, /marketplace, /marketplace/{slug}, /mleq", "Loading buffers appeared first and cleared into real content once data resolved; no console errors were reported."),
        ],
        [1900, 3300, 4160],
        font_size=8.5,
    )

    add_heading(doc, "9. Deployment and Environment Notes", 1)
    add_bullet(doc, "Development runs with next dev on localhost:3001.")
    add_bullet(doc, "The dashboard API uses environment variables for registry owner, repository, branch, default model, upstream API base URL, account starting capital, and GitHub read tokens.")
    add_bullet(doc, "When running on localhost or 127.0.0.1, the dashboard route prefers upstream QSentia API data unless disabled by QSENTIA_DISABLE_UPSTREAM_API_PROXY=1.")
    add_bullet(doc, "Server-side GitHub token candidates include GITHUB_READ_TOKEN, QSENTIA_GITHUB_READ_TOKEN, GITHUB_TOKEN, GH_TOKEN, VERCEL_GITHUB_TOKEN, and NEXT_PUBLIC_GITHUB_READ_TOKEN.")
    add_bullet(doc, "The local last-good cache is stored in .qsentia-cache/dashboard-last-good.json and should be treated as operational cache, not source of record.")

    add_heading(doc, "10. Recommended Next Steps", 1)
    add_numbered(doc, "Create formal requirement IDs for every page and API endpoint, then connect them to acceptance tests.")
    add_numbered(doc, "Add automated API contract tests for /api/dashboard, /api/models, /api/models/{slug}, and /api/models/{slug}/demo.")
    add_numbered(doc, "Add monitored cache-age and upstream-health status so stale data is visible to operators.")
    add_numbered(doc, "Replace local-storage-only sign-in with production authentication before treating accounts as secure.")
    add_numbered(doc, "Archive or remove legacy *-SHLOK files after confirming they are no longer part of the active implementation.")
    add_numbered(doc, "Create a user-facing glossary for normalized equity index, benchmark, source points, drawdown, Sharpe ratio, and execution audit.")

    add_heading(doc, "Appendix A. Current File Map", 1)
    add_para(doc, "Active implementation files are concentrated in app/, components/, and lib/. The following inventory summarizes the main wiring files.")
    add_table(
        doc,
        ["File", "Role"],
        [
            ("app/layout.tsx", "Global metadata, footer, and Alex assistant injection."),
            ("app/page.tsx", "Home page, selected model equity index, API preview, mission/vision/objective section."),
            ("app/dashboard/page.tsx", "Live dashboard terminal and model selector."),
            ("components/ResearchTerminal.tsx", "Research tabs, tickers, filters, normalized charts, exposure, execution, decision, and health views."),
            ("app/mleq/page.tsx", "Machine Learning Equity Quant narrative and live telemetry page."),
            ("app/marketplace/page.tsx", "API-backed marketplace model list."),
            ("app/marketplace/[slug]/page.tsx", "API-backed model detail page and demo preview button."),
            ("app/api/dashboard/route.ts", "Central dashboard API orchestration, upstream proxy, cache, GitHub/raw log fallback, metrics and benchmark computation."),
            ("app/api/models/route.ts", "Marketplace model list API."),
            ("app/api/models/[slug]/route.ts", "Single model detail API."),
            ("app/api/models/[slug]/demo/route.ts", "Rate-limited latest decision preview API."),
            ("components/PageChrome.tsx", "Site header, page shell, section cards, empty states, technical backdrop, and shared ApiLoadingPanel."),
            ("components/AlexAssistant.tsx", "Global help assistant with session personalization and live API awareness."),
            ("lib/metrics.ts", "Return, normalized index, drawdown, stats, and formatting helpers."),
            ("lib/modelCatalog.ts", "Dashboard-to-marketplace model mapping and model detail/preview helpers."),
            ("lib/clientSession.ts", "Client local storage session utilities."),
        ],
        [3600, 5760],
        font_size=8.2,
    )

    add_heading(doc, "Appendix B. Data Dictionary", 1)
    add_table(
        doc,
        ["Term", "Meaning in this site"],
        [
            ("Model equity index", "Normalized strategy portfolio track record where the first observation equals 100."),
            ("Current index", "Latest normalized portfolio value. For example, 102.29 indicates a 2.29% gain over the displayed window."),
            ("Source points", "Number of source observations used for the selected chart or metric."),
            ("Benchmark series", "SPY, QQQ, DIA, IWM, and VTI comparison data aligned to model windows."),
            ("Execution audit", "Decision, submitted order, planned order, position, target-weight, and signal-history rows exposed for review."),
            ("Live paper active", "Paper execution state inferred from positions/orders/health data."),
            ("Last-good cache", "Local cached dashboard payload used when upstream data is temporarily unavailable."),
        ],
        [2300, 7060],
        font_size=8.8,
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build_report()
    print(path)
